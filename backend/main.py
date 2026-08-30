import os
from contextlib import asynccontextmanager
import asyncio
import threading
from fastapi import Depends, FastAPI, HTTPException, UploadFile, File, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, StreamingResponse
from pydantic import BaseModel
from datetime import datetime, timezone
from io import BytesIO
import logging
import tempfile
import time
from typing import Optional
import uuid

from backend.database import engine, Base, SessionLocal, using_ephemeral_sqlite
from backend.routers import legal_routes
from backend.routers import history_routes
from backend.routers import obligations_routes
from backend.routers.notifications import router as notifications_router
from backend.routers.compare_routes import router as compare_router
from backend.routers import export_routes
from backend.routers.collaboration_routes import router as collaboration_router
from backend.routers.comments_routes import router as comments_router
from backend.routers import feedback_routes
from backend.routers.developer_routes import router as developer_router
from backend.auth import validate_token_or_api_key, AuthIdentity
from backend.utils.limiter import create_rate_limiter, SimpleRateLimiter
from backend.utils.cleanup import start_token_cleanup_task
from backend.services.reminder_service import run_obligation_reminders
from backend.config import get_settings
from apscheduler.schedulers.asyncio import AsyncIOScheduler
from backend.storage.upload_tasks import get_upload_task_storage
from backend.services.upload_job_queue import (
    UploadJobQueue,
    build_upload_job,
    process_upload_job_async,
)

# Optional imports (wrap in try/except so server can start without optional deps)
try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None

try:
    from docx import Document as DocxDocument  # type: ignore[import-untyped]
except Exception:
    DocxDocument = None  # type: ignore[assignment,misc]

# Import pipeline exceptions, validations, and service
from backend.core.exceptions import (
    AIError, ValidationError, ProviderError, TimeoutError, ServiceUnavailableError
)
from backend.core.validation import (
    validate_chat_input, validate_summarize_input, validate_simplify_input, sanitize_text, validate_mime_and_bytes,
    validate_docx_archive_safety, validate_jurisdiction
)
from backend.services.ai_service import ai_service, correlation_id_var
from backend.services.rag_service import rag_service
from backend.services.cache_service import semantic_cache

#Middleware import 
from backend.middleware.rate_limit import RateLimitMiddleware
from backend.middleware.correlation_id import validate_or_generate_correlation_id
# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Get configuration from centralized settings
settings = get_settings()
file_config = settings.file_upload
rate_config = settings.rate_limit
cors_config = settings.cors

# Track application start time for uptime calculation
_app_start_time = time.monotonic()

@asynccontextmanager
async def lifespan(app: FastAPI):
    # Vercel may create and freeze function instances between requests. Long-
    # lived cleanup/scheduler tasks do not have a reliable lifecycle there and
    # can keep a serverless invocation alive unnecessarily.
    if os.getenv("VERCEL"):
        yield
        return

    # Start the token blacklist cleanup worker in the background
    cleanup_interval = file_config.token_cleanup_interval_seconds
    cleanup_task = asyncio.create_task(start_token_cleanup_task(interval_seconds=cleanup_interval))

    # Daily obligation-reminder job (30/15/1-day-out thresholds).
    scheduler = AsyncIOScheduler()
    scheduler.add_job(run_obligation_reminders, "interval", hours=24, id="obligation_reminders")
    scheduler.start()

    try:
        yield
    finally:
        scheduler.shutdown(wait=False)
        cleanup_task.cancel()
        try:
            await cleanup_task
        except asyncio.CancelledError:
            pass

app = FastAPI(lifespan=lifespan)


# Exception Handlers to match standardized error contracts
@app.exception_handler(ValidationError)
async def validation_exception_handler(request: Request, exc: ValidationError):
    logger.warning(f"[{correlation_id_var.get()}] Validation error: {exc}")
    return JSONResponse(
        status_code=400,
        content={
            "error": "validation_error",
            "detail": str(exc),
            "correlation_id": correlation_id_var.get()
        }
    )


@app.exception_handler(ProviderError)
async def provider_exception_handler(request: Request, exc: ProviderError):
    logger.error(f"[{correlation_id_var.get()}] Upstream provider error: {exc}")
    return JSONResponse(
        status_code=502,
        content={
            "error": "provider_error",
            "detail": str(exc),
            "correlation_id": correlation_id_var.get()
        }
    )


@app.exception_handler(TimeoutError)
async def timeout_exception_handler(request: Request, exc: TimeoutError):
    logger.error(f"[{correlation_id_var.get()}] Request timeout: {exc}")
    return JSONResponse(
        status_code=504,
        content={
            "error": "timeout_error",
            "detail": str(exc),
            "correlation_id": correlation_id_var.get()
        }
    )


@app.exception_handler(ServiceUnavailableError)
async def service_unavailable_exception_handler(request: Request, exc: ServiceUnavailableError):
    logger.error(f"[{correlation_id_var.get()}] Service unavailable: {exc}")
    return JSONResponse(
        status_code=503,
        content={
            "error": "service_unavailable",
            "detail": str(exc),
            "correlation_id": correlation_id_var.get()
        }
    )

import sys

# Global unhandled HTTP exceptions
@app.exception_handler(Exception)
async def global_exception_handler(request: Request, exc: Exception):
    corr_id = correlation_id_var.get()
    logger.error(f"[{corr_id}] Unhandled exception: {exc}", exc_info=True)
    return JSONResponse(
        status_code=500,
        content={
            "error": "internal_server_error",
            "detail": "An unexpected error occurred.",
            "correlation_id": corr_id
        }
    )

# Global unhandled thread/process exceptions
def handle_uncaught_exception(exc_type, exc_value, exc_traceback):
    if not issubclass(exc_type, Exception):
        sys.__excepthook__(exc_type, exc_value, exc_traceback)
        return
    logger.critical("Uncaught global exception", exc_info=(exc_type, exc_value, exc_traceback))

sys.excepthook = handle_uncaught_exception


# Create database tables once per process. A database outage should be
# reported by /health and by the affected endpoint, not turn every Vercel
# invocation into an opaque FUNCTION_INVOCATION_FAILED response.
database_initialization_error: Optional[str] = None
try:
    Base.metadata.create_all(bind=engine)
except Exception as exc:
    database_initialization_error = type(exc).__name__
    logger.exception(
        "Database schema initialization failed; requests will report database unavailability."
    )

# Authentication removed — LegalEase runs as a single local/anonymous user.
# Include legal mapping router
app.include_router(legal_routes.router)
# Include notifications router
app.include_router(notifications_router)
# Include obligations router
app.include_router(obligations_routes.router)
# Include history router
app.include_router(history_routes.router)
# Include multi-document comparison router
app.include_router(compare_router)
# Include export router
app.include_router(export_routes.router)
app.include_router(collaboration_router)
app.include_router(comments_router)
app.include_router(feedback_routes.router)
app.include_router(developer_router)


# Environment configuration - defaults to production for security
ENVIRONMENT = os.getenv("ENVIRONMENT", "production").lower()

configured_allowed_origins = os.getenv("ALLOWED_ORIGINS")
configured_frontend_url = os.getenv("FRONTEND_URL")
if configured_allowed_origins is None and configured_frontend_url:
    # FRONTEND_URL is a documented fallback when ALLOWED_ORIGINS is omitted.
    raw_allowed_origins = configured_frontend_url
else:
    raw_allowed_origins = cors_config.allowed_origins or cors_config.frontend_url
ALLOWED_ORIGINS = [
    origin.strip().rstrip("/")
    for origin in raw_allowed_origins.split(",")
    if origin.strip()
]

# Add localhost origins only if explicitly configured AND not in production
if cors_config.allow_localhost_cors and ENVIRONMENT in ("development", "testing", "local"):
    localhost_origins = []
    for host in ["http://localhost", "http://127.0.0.1"]:
        for port in range(5173, 5181):
            localhost_origins.append(f"{host}:{port}")
    
    # Add localhost origins that aren't already in the list
    added_count = 0
    for origin in localhost_origins:
        if origin not in ALLOWED_ORIGINS:
            ALLOWED_ORIGINS.append(origin)
            added_count += 1
    
    if added_count > 0:
        logger.info(f"Localhost CORS enabled: added {added_count} localhost development origins")

# Validate and deduplicate origins
validated_origins = []
seen_origins = set()
for origin in ALLOWED_ORIGINS:
    # Basic origin format validation
    if not origin:
        continue
    
    # Check if it looks like a valid origin (scheme://host or scheme://host:port)
    if not (origin.startswith("http://") or origin.startswith("https://")):
        logger.warning(f"Invalid CORS origin ignored (missing scheme): {origin}")
        continue
    
    # Remove duplicates while preserving order
    if origin not in seen_origins:
        seen_origins.add(origin)
        validated_origins.append(origin)

ALLOWED_ORIGINS = validated_origins
# Rate-limit middleware registered first so that CORSMiddleware
# (added second) wraps it — ensuring 429 responses include CORS headers.
app.add_middleware(RateLimitMiddleware)
app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["X-Citations"],
)
logger.info(f"Allowed frontend origins: {ALLOWED_ORIGINS}")


# Correlation ID middleware to inject trace headers
@app.middleware("http")
async def correlation_id_middleware(request: Request, call_next):
    # Validate or generate safe correlation ID
    client_id = request.headers.get("X-Correlation-ID")
    corr_id, was_valid = validate_or_generate_correlation_id(client_id)
    
    token = correlation_id_var.set(corr_id)
    try:
        response = await call_next(request)
        response.headers["X-Correlation-ID"] = corr_id
        return response
    finally:
        correlation_id_var.reset(token)


# Configuration
MAX_UPLOAD_SIZE = file_config.max_upload_size
CHUNK_SIZE = 1024 * 1024
MAX_PDF_PAGES = file_config.max_pdf_pages
MAX_DOCX_PARAGRAPHS = file_config.max_docx_paragraphs
MAX_EXTRACTED_TEXT_CHARS = file_config.max_extracted_text_chars
UPLOAD_PARSE_TIMEOUT_SECONDS = file_config.upload_parse_timeout_seconds


RATE_LIMIT_PERIOD = rate_config.rate_limit_period
RATE_LIMIT_KEY_CALLS = rate_config.rate_limit_key_calls


# Defaults: 300 requests per minute per API key
key_limiter = create_rate_limiter(calls=RATE_LIMIT_KEY_CALLS, period=RATE_LIMIT_PERIOD)


class ChatRequest(BaseModel):
    message: str
    context: Optional[str] = None
    conversation_history: Optional[list[dict[str, str]]] = None
    stream: Optional[bool] = False
    jurisdiction: str = "General / Not Specified"


class EditMessageRequest(BaseModel):
    """Request body for the message-edit/branching endpoint (#366)."""
    new_content: str
    conversation_history: Optional[list[dict[str, str]]] = None
    context: Optional[str] = None


class SummarizeRequest(BaseModel):
    text: str


class TLDRRequest(BaseModel):
    text: str


class TLDRResponse(BaseModel):
    parties: str = "Not specified"
    deadlines: str = "Not specified"
    financials: str = "Not specified"
    penalties: str = "Not specified"
    key_takeaways: list[str] = ["Not specified"]


class SimplifyRequest(BaseModel):
    text: str


class SimplifyResponse(BaseModel):
    simplifiedText: str


class HealthResponse(BaseModel):
    status: str
    uptime_seconds: float
    timestamp: str
    details: Optional[dict] = None


def _get_client_ip(request: Request) -> str:
    forwarded_for = request.headers.get("x-forwarded-for")
    if forwarded_for:
        return forwarded_for.split(",", 1)[0].strip()
    return request.client.host if request.client else "unknown"


def _extract_pdf_text(file_path: str) -> str:
    if fitz is None:
        raise HTTPException(status_code=503, detail="PDF processing not available")

    doc = None
    try:
        doc = fitz.open(file_path)
        if doc.page_count > MAX_PDF_PAGES:
            raise HTTPException(status_code=413, detail="PDF is too large to process safely")

        extracted_parts = []
        for page in doc:
            extracted_parts.append(page.get_text())
            if len("".join(extracted_parts)) >= MAX_EXTRACTED_TEXT_CHARS:
                break

        return "".join(extracted_parts)
    finally:
        if doc is not None:
            doc.close()


def _extract_docx_text(file_path: str) -> str:
    if DocxDocument is None:
        raise HTTPException(status_code=503, detail="DOCX processing not available")

    document = None
    try:
        document = DocxDocument(file_path)
        if len(document.paragraphs) > MAX_DOCX_PARAGRAPHS:
            raise HTTPException(status_code=413, detail="DOCX is too large to process safely")

        return "\n".join(
            paragraph.text
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        )
    finally:
        close_method = getattr(document, "close", None) if document is not None else None
        if callable(close_method):
            close_method()


async def _run_bounded_parser(parser, file_path: str) -> str:
    try:
        return await asyncio.wait_for(
            asyncio.to_thread(parser, file_path),
            timeout=UPLOAD_PARSE_TIMEOUT_SECONDS,
        )
    except asyncio.TimeoutError:
        raise HTTPException(status_code=413, detail="File is too complex to process safely")


@app.post("/chat")
async def chat(request: Request, payload: ChatRequest, identity: AuthIdentity = Depends(validate_token_or_api_key)):
    # Rate limiting using the authenticated identity
    if not key_limiter.check(identity.get_rate_limit_key())["allowed"]:
        raise HTTPException(status_code=429, detail="Rate limit exceeded")

    # Sanitize inputs
    sanitized_message = sanitize_text(payload.message)
    sanitized_context = sanitize_text(payload.context) if payload.context else None
    citations = []

    # Handle RAG context retrieval for non-streaming requests early.
    # Streaming requests handle RAG inside the generator to avoid blocking initial response.
    if sanitized_context and not payload.stream:
        import hashlib
        
        doc_hash = hashlib.md5(sanitized_context.encode()).hexdigest()
        try:
            if doc_hash not in rag_service.indexed_docs:
                await rag_service.add_document(sanitized_context, doc_hash)
            sanitized_context, citations = await rag_service.get_context(sanitized_message, doc_hash)
        except Exception as e:
            logger.warning(f"RAG retrieval failed: {e}. Falling back to non-RAG heuristic.")
            # Fall back gracefully to a non-RAG heuristic (truncating context)
            if len(sanitized_context) > 5000:
                sanitized_context = sanitized_context[:5000] + "\n... [Context truncated due to RAG failure] ..."

    # Early payload validation
    validate_chat_input(sanitized_message, sanitized_context)
    validate_jurisdiction(payload.jurisdiction)

    cache_key = f"{sanitized_message} || {sanitized_context}" if sanitized_context else sanitized_message
    cache_namespace = identity.get_rate_limit_key()
    cached_response = semantic_cache.get(cache_key, namespace=cache_namespace)

    if cached_response:
        logger.info(f"[{correlation_id_var.get()}] Serving response from semantic cache")
        if payload.stream:
            async def stream_cached():
                import asyncio
                import json
                words = cached_response.split(" ")
                for i, word in enumerate(words):
                    chunk = word + (" " if i < len(words) - 1 else "")
                    payload_data = json.dumps({"response": chunk})
                    yield f"data: {payload_data}\n\n"
                    await asyncio.sleep(0.01)
                yield "data: [DONE]\n\n"
            return StreamingResponse(stream_cached(), media_type="text/event-stream")
        else:
            return {"response": cached_response}

    # Streaming or standard block handling
    if payload.stream:
        async def stream_generator():
            nonlocal sanitized_context, citations
            full_response = ""
            try:
                # Perform RAG retrieval inside the stream generator asynchronously to prevent blocking the initial response
                if sanitized_context:
                    import hashlib
                    
                    doc_hash = hashlib.md5(sanitized_context.encode()).hexdigest()
                    try:
                        if doc_hash not in rag_service.indexed_docs:
                            await rag_service.add_document(sanitized_context, doc_hash)
                        sanitized_context, citations = await rag_service.get_context(sanitized_message, doc_hash)
                    except Exception as e:
                        logger.warning(f"RAG retrieval failed inside stream generator: {e}. Falling back to non-RAG heuristic.")
                        if len(sanitized_context) > 5000:
                            sanitized_context = sanitized_context[:5000] + "\n... [Context truncated due to RAG failure] ..."

                async for chunk in ai_service.generate_chat_response(
                    message=sanitized_message,
                    context=sanitized_context,
                    history=payload.conversation_history,
                    stream=True,
                    jurisdiction=payload.jurisdiction
                ):
                    # Clean the SSE format to cache clean raw text
                    if chunk.startswith("data: ") and chunk != "data: [DONE]\n\n":
                        try:
                            import json
                            json_str = chunk.replace("data: ", "").strip()
                            data = json.loads(json_str)
                            full_response += data.get("response", "")
                        except Exception:
                            pass
                    yield chunk
                
                # Re-calculate the cache key using the post-RAG context
                final_cache_key = f"{sanitized_message} || {sanitized_context}" if sanitized_context else sanitized_message
                semantic_cache.set(final_cache_key, full_response, namespace=cache_namespace)
            except Exception as e:
                logger.error(f"[{correlation_id_var.get()}] Stream generation error: {e}")
                yield "\n[Error: Inference stream failed]"

        headers = {}
        if citations:
            import json, base64
            headers["X-Citations"] = base64.b64encode(json.dumps(citations).encode()).decode()
        return StreamingResponse(stream_generator(), media_type="text/event-stream", headers=headers)
    else:
        response_gen = ai_service.generate_chat_response(
            message=sanitized_message,
            context=sanitized_context,
            history=payload.conversation_history,
            stream=False,
            jurisdiction=payload.jurisdiction
        )
        response_text = ""
        async for chunk in response_gen:
            response_text += chunk
        semantic_cache.set(cache_key, response_text, namespace=cache_namespace)
        return {"response": response_text, "citations": citations}


@app.post("/upload", status_code=202)
async def upload_document(
    request: Request,
    file: UploadFile = File(...),
    identity: AuthIdentity = Depends(validate_token_or_api_key)
):
    """Accept a document upload, immediately return 202 with a task_id,
    and enqueue it for durable background processing (#365)."""
    # Content-Length pre-check
    try:
        content_length = int(request.headers.get("content-length", "0"))
    except Exception:
        content_length = 0
    if content_length and content_length > MAX_UPLOAD_SIZE:
        raise HTTPException(status_code=413, detail="Uploaded file is too large")

    task_id = str(uuid.uuid4())
    temp_path = None
    try:
        filename = file.filename or "unknown"
        file_extension = os.path.splitext(filename)[1].lower()
        total_size = 0
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension or "") as temp_file:
            temp_path = temp_file.name
            while True:
                chunk = await file.read(CHUNK_SIZE)
                if not chunk:
                    break
                total_size += len(chunk)
                if total_size > MAX_UPLOAD_SIZE:
                    raise HTTPException(status_code=413, detail="Uploaded file is too large")
                temp_file.write(chunk)

        # Read only first 4096 bytes for MIME validation
        with open(temp_path, "rb") as temp_file:
            content_prefix = temp_file.read(4096)

        validate_mime_and_bytes(content_prefix, file.content_type or "", filename)

        if file_extension == ".docx":
            validate_docx_archive_safety(temp_path)

    except ValidationError as e:
        if temp_path and os.path.exists(temp_path):
            os.unlink(temp_path)
        raise HTTPException(status_code=400, detail=str(e))
    except HTTPException:
        if temp_path and os.path.exists(temp_path):
            os.unlink(temp_path)
        raise
    except Exception as e:
        if temp_path and os.path.exists(temp_path):
            os.unlink(temp_path)
        logger.error(f"Upload error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail="Failed to receive document")

    job = build_upload_job(
        task_id=task_id,
        file_path=temp_path,
        filename=filename,
        content_type=file.content_type or "",
        file_extension=file_extension,
        content_prefix=content_prefix,
    )

    # Vercel does not keep a worker process alive after the request returns.
    # Require shared task storage and finish extraction inside the invocation
    # so accepted uploads cannot remain queued forever.
    if os.getenv("VERCEL"):
        task_storage = get_upload_task_storage()
        if not task_storage.using_redis:
            if temp_path and os.path.exists(temp_path):
                os.unlink(temp_path)
            raise HTTPException(
                status_code=503,
                detail="Document processing requires REDIS_URL on Vercel",
            )

        task_storage.create_task(
            task_id,
            status="processing",
            progress=0,
            result=None,
            user_id=identity.get_user_id(),
        )
        try:
            await process_upload_job_async(job)
        except Exception:
            # The worker helper records the failure in Redis. Return the task
            # so the client can display its normal failed-processing state.
            logger.exception("[%s] Inline Vercel upload processing failed", task_id)

        task = task_storage.get_task(task_id) or {}
        return JSONResponse(
            status_code=202,
            content={
                "task_id": task_id,
                "filename": filename,
                "status": task.get("status", "failed"),
            },
        )

    # Register the task as "queued" and launch the background worker.
    task_storage = get_upload_task_storage()
    task_storage.create_task(
        task_id,
        status="queued",
        progress=0,
        result=None,
        user_id=identity.get_user_id(),
    )
    logger.info(f"[{task_id}] Upload task created")

    try:
        job_queue = UploadJobQueue()
    except RuntimeError:
        task_storage.mark_failed(task_id, "Document processing is not configured")
        if temp_path and os.path.exists(temp_path):
            os.unlink(temp_path)
        raise HTTPException(
            status_code=503,
            detail="Document processing requires REDIS_URL in production",
        )

    if not job_queue.enqueue(job):
        task_storage.mark_failed(task_id, "Failed to enqueue upload job")
        if temp_path and os.path.exists(temp_path):
            os.unlink(temp_path)
        raise HTTPException(status_code=503, detail="Failed to enqueue document for processing")

    # In non-production/local test runs, kick off an in-process worker so the
    # task moves past the initial queued state without requiring a separate
    # worker process. Durable Redis-backed deployments still rely on the
    # external worker loop.
    if not job_queue.using_redis and settings.environment.environment in {"development", "testing", "local"}:
        threading.Thread(
            target=lambda: asyncio.run(process_upload_job_async(job)),
            daemon=True,
        ).start()

    return JSONResponse(
        status_code=202,
        content={"task_id": task_id, "filename": filename, "status": "queued"},
    )


@app.put("/chat/messages/{message_id}")
async def edit_message(
    message_id: str,
    payload: EditMessageRequest,
    request: Request,
    identity: AuthIdentity = Depends(validate_token_or_api_key),
):
    """Edit a previous user message and regenerate the AI response (#366).

    The frontend passes the edited text plus the conversation history up to
    (but not including) the message being edited.  The backend re-runs the AI
    and returns a fresh assistant response, leaving branching bookkeeping to
    the client-side storage layer.
    """
    sanitized_message = sanitize_text(payload.new_content)
    sanitized_context = sanitize_text(payload.context) if payload.context else None
    validate_chat_input(sanitized_message, sanitized_context)

    response_gen = ai_service.generate_chat_response(
        message=sanitized_message,
        context=sanitized_context,
        history=payload.conversation_history,
        stream=False,
    )
    response_text = ""
    async for chunk in response_gen:
        response_text += chunk

    return {
        "message_id": message_id,
        "edited_content": sanitized_message,
        "response": response_text,
    }


@app.get("/upload/status/{task_id}")
async def upload_status(task_id: str, identity: AuthIdentity = Depends(validate_token_or_api_key)):
    """
    Poll the processing status of an async upload task (#365).

    Returns 404 (not 403) when the task exists but belongs to a different
    identity, matching the ownership-check pattern used elsewhere (e.g.
    comments_routes.py's _get_owned_document) — this avoids confirming to
    an unauthorized caller that a given task_id actually exists.
    """
    task_storage = get_upload_task_storage()
    task = task_storage.get_task(task_id)
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")

    if task.get("user_id") != identity.get_user_id():
        raise HTTPException(status_code=404, detail="Task not found")

    response = {
        "task_id": task_id,
        "status": task["status"],
        "progress": task["progress"],
        "result": task["result"],
    }
    return response



@app.post("/summarize")
async def summarize(request: Request, payload: SummarizeRequest, identity: AuthIdentity = Depends(validate_token_or_api_key)):
    # Rate limiting using the authenticated identity (already applied to
    # /chat and /simplify; this endpoint was missing it)
    if not key_limiter.check(identity.get_rate_limit_key())["allowed"]:
        raise HTTPException(status_code=429, detail="Rate limit exceeded")

    # Sanitize input
    sanitized_text = sanitize_text(payload.text)

    # Early payload validation
    validate_summarize_input(sanitized_text)

    summary = await ai_service.generate_summary(sanitized_text)
    return {"summary": summary}


@app.post("/api/tldr", response_model=TLDRResponse)
@app.post("/tldr", response_model=TLDRResponse)
async def tldr(request: Request, payload: TLDRRequest, identity: AuthIdentity = Depends(validate_token_or_api_key)):
    if not key_limiter.check(identity.get_rate_limit_key())["allowed"]:
        raise HTTPException(status_code=429, detail="Rate limit exceeded")

    sanitized_text = sanitize_text(payload.text)
    validate_summarize_input(sanitized_text)

    tldr_data = await ai_service.generate_tldr(sanitized_text)
    return TLDRResponse(**tldr_data)


@app.post("/api/simplify", response_model=SimplifyResponse)
@app.post("/simplify", response_model=SimplifyResponse)
async def simplify(request: Request, payload: SimplifyRequest, identity: AuthIdentity = Depends(validate_token_or_api_key)):
    # Rate limiting using the authenticated identity
    if not key_limiter.check(identity.get_rate_limit_key())["allowed"]:
        raise HTTPException(status_code=429, detail="Rate limit exceeded")

    # Sanitize input
    sanitized_text = sanitize_text(payload.text)

    # Early payload validation
    validate_simplify_input(sanitized_text)

    # Log the action
    logger.info(f"[{correlation_id_var.get()}] Simplifying clause/text of length {len(sanitized_text)}")

    simplified = await ai_service.simplify_clause(sanitized_text)
    return SimplifyResponse(simplifiedText=simplified)


@app.get("/health", response_model=HealthResponse)
async def health():
    """
    Health check endpoint with structured response.
    Returns HTTP 503 when the service is degraded.
    """
    health_data = ai_service.check_health()
    rag_health = rag_service.check_health()
    uptime = time.monotonic() - _app_start_time
    timestamp = datetime.now(timezone.utc).replace(tzinfo=None).isoformat() + "Z"

    # Database connectivity check
    db_status = "ok"
    try:
        from sqlalchemy import text
        from backend.database import SessionLocal as HealthSessionLocal
        db = HealthSessionLocal()
        db.execute(text("SELECT 1"))
        db.close()
    except Exception as e:
        logger.error(f"Database health check failed: {e}")
        db_status = "down"

    if using_ephemeral_sqlite or database_initialization_error:
        db_status = "down"

    redis_status = "ok"
    redis_backend = getattr(key_limiter, "_redis_backend", None)
    if redis_backend:
        try:
            h_res = redis_backend.health_check()
            if not h_res.get("healthy", True):
                redis_status = "degraded"
        except Exception as e:
            logger.error(f"Redis health check failed: {e}")
            redis_status = "down"
    elif ENVIRONMENT in {"production", "staging"}:
        # The fallback is process-local and cannot support durable upload task
        # state or distributed rate limits on a serverless deployment.
        redis_status = "not_configured"

    status = health_data.get("status", "unknown")
    if rag_health.get("status") in {"degraded", "failed"} and status == "ok":
        status = "degraded"
    if db_status == "down":
        status = "degraded"
    if redis_status in ("degraded", "down", "not_configured"):
        status = "degraded"

    details = health_data.get("details") or {}
    if not isinstance(details, dict):
        details = {"ai_details": details}
    details["rag"] = rag_health
    details["database"] = db_status
    details["database_persistence"] = "ephemeral" if using_ephemeral_sqlite else "persistent"
    details["redis"] = redis_status

    response = HealthResponse(
        status=status,
        uptime_seconds=round(uptime, 2),
        timestamp=timestamp,
        details=details,
    )

    if response.status == "degraded":
        return JSONResponse(status_code=503, content={"detail": response.model_dump()})

    return response


if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="0.0.0.0", port=8000)
 
